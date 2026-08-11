from pathlib import Path
from docx import Document

for path in sorted(Path("docs/26FA_v0.9").glob("*.docx")):
    doc = Document(path)
    print(f"\n=== {path.name} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)} ===")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().replace("\n", " / ")
        if text:
            print(f"P{i} [{p.style.name}] {text[:260]}")
    for ti, table in enumerate(doc.tables):
        print(f"TABLE {ti} {len(table.rows)}x{len(table.columns)}")
        for row in table.rows[:12]:
            print(" | ".join(c.text.strip().replace("\n", " / ")[:140] for c in row.cells))
