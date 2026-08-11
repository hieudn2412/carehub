from pathlib import Path

from docx import Document


FILES = [
    Path("docs/docs-old/Report 1 - Project Introduction (4).docx"),
    Path("docs/docs-old/Report 2.0_Project_Plan_VietDuc-Care_v1.0.0 (1).docx"),
    Path("docs/docs-old/Report 3.0_SRS_VietDuc-Care_v1.6.docx"),
    Path("docs/docs-old/Report 3.2_FDS_VietDuc-Care_v1.1.0.docx"),
    Path("docs/docs-old/Report 4_TDS_VietDuc-Care_v1.1.0.docx"),
    Path("docs/docs-old/VietDuc-Care_TestPlan.docx"),
]


for path in FILES:
    document = Document(path)
    print(f"\nFILE {path.name} paragraphs={len(document.paragraphs)} tables={len(document.tables)}")
    for paragraph in document.paragraphs:
        style = paragraph.style.name or ""
        text = paragraph.text.strip()
        if text and (style.startswith("Heading") or style in {"Title", "Subtitle"}):
            print(style, repr(text[:240]))
    if path.name.startswith("Report 1") or "Project_Plan" in path.name:
        for index, table in enumerate(document.tables[:20], start=1):
            print(f"TABLE {index}")
            for row in table.rows:
                print(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells))
