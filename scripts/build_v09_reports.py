import hashlib
import json
import re
import shutil
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATES = ROOT / "docs" / "26FA_v0.9"
OUT = ROOT / "docs" / "26FA_v0.9_CareHub"
ARTIFACTS = ROOT / "report-work" / "artifacts"
DIAGRAMS = ROOT / "report-work" / "diagrams"
DATE = "11/08/2026"
VERSION = "v0.9.0"
PROJECT = "CareHub – Viet Duc Hospital Nursing Quality and Competency Management System"
AUTHORS = "Đỗ Mạnh Tuấn; Dương Ngọc Hiếu; Ngô Trung Kiên; Phí Hoàng Nam; Phạm Hải Nam"


FEATURES = [
    ("FT-01", "Authentication & Account Management", "Must", "Access", "UC-01–UC-03", "BF-01", "Secure sign-in, token refresh, password recovery, profile and administrative account lifecycle."),
    ("FT-02", "System Configuration & Logging", "Must", "Administration", "UC-04–UC-05", "BF-02", "Hospital-wide settings, import history, evaluation audit trail and operational logs."),
    ("FT-03", "Reference Data Management", "Must", "Administration", "UC-06–UC-08", "BF-02/BF-12", "Departments, positions, education levels, employees and controlled imports."),
    ("FT-04", "Training Records Management", "Must", "Training", "UC-09–UC-10", "BF-03", "Create, edit, submit and cancel CME records with evidence and audit history."),
    ("FT-05", "Training Compliance Tracking", "Must", "Training", "UC-11–UC-13", "BF-04", "Calculate submitted hours against the global training target and expose scoped employee ledgers."),
    ("FT-06", "Training Monitoring & Alerting", "Should", "Training", "UC-14", "BF-04/BF-10", "Scan compliance state and publish timely in-system or email notifications."),
    ("FT-07", "Training Analytics Dashboard & Export", "Should", "Analytics", "UC-15", "BF-11", "Training summaries, department comparison, trends and controlled exports."),
    ("FT-08", "Quality Form Management", "Must", "Quality", "UC-16–UC-18", "BF-05", "Versioned form/checklist design, assignment, response capture and history."),
    ("FT-09", "Quality Scoring Engine", "Must", "Quality", "UC-17–UC-19", "BF-06", "Configurable scoring, critical weights, pass thresholds and recalculation."),
    ("FT-10", "Quality Benchmark & Alerting", "Should", "Quality", "UC-20–UC-21", "BF-06/BF-10", "Hospital/department targets and below-target alerting."),
    ("FT-11", "Quality Analytics Dashboard & Export", "Should", "Analytics", "UC-22", "BF-11", "Checklist performance, trends, compliance and export views."),
    ("FT-12", "Test & Assessment Management", "Must", "Evaluation", "UC-23–UC-25", "BF-07", "Question bank, sets, categories, exam configuration, papers and assignments."),
    ("FT-13", "Test Execution & Tracking", "Must", "Evaluation", "UC-26–UC-27", "BF-08", "Attempt lifecycle, timer, answer capture, submission, grading and history."),
    ("FT-14", "Performance Classification Engine", "Must", "Evaluation", "UC-28", "BF-08", "Classify competency on the implemented 0–10 score scale using configurable rules."),
    ("FT-15", "Performance Monitoring Dashboard & Export", "Should", "Analytics", "UC-29", "BF-11", "Competency, field, technique and department-level summaries."),
    ("FT-16", "Notification & Communication Management", "Should", "Notification", "UC-30–UC-31", "BF-10", "Notification inbox, policies, email templates and RabbitMQ delivery."),
    ("FT-17", "AI-Assisted Question Generation", "Could", "AI Evaluation", "UC-32–UC-34", "BF-09", "Document ingestion, DeepSeek generation, E5 semantic deduplication, VietQuill paraphrase and reviewer approval."),
]

UCS = [
    ("UC-01", "Manage Authentication", "Guest / all authenticated roles", "FT-01", "BF-01"),
    ("UC-02", "Manage User Profile", "USER, MANAGER, ADMIN", "FT-01", "BF-01"),
    ("UC-03", "Manage Account", "ADMIN", "FT-01", "BF-01"),
    ("UC-04", "Manage System Configuration", "ADMIN", "FT-02", "BF-02"),
    ("UC-05", "Manage System & Import Logs", "ADMIN", "FT-02", "BF-02/BF-12"),
    ("UC-06", "Manage Employee Reference Data", "ADMIN", "FT-03", "BF-02"),
    ("UC-07", "Manage Department Reference Data", "ADMIN", "FT-03", "BF-02"),
    ("UC-08", "Import Reference Data", "ADMIN", "FT-03", "BF-12"),
    ("UC-09", "Manage Training Records", "USER, MANAGER, ADMIN", "FT-04", "BF-03"),
    ("UC-10", "Manage Training Activity Types", "ADMIN", "FT-04", "BF-03"),
    ("UC-11", "Configure Training Requirements", "ADMIN", "FT-05", "BF-04"),
    ("UC-12", "Track Training Compliance", "USER, MANAGER, ADMIN", "FT-05", "BF-04"),
    ("UC-13", "Review Training Evidence", "MANAGER, ADMIN", "FT-05", "BF-03"),
    ("UC-14", "Manage Training Alerts", "ADMIN", "FT-06", "BF-04/BF-10"),
    ("UC-15", "Analyze Training Dashboard & Reports", "MANAGER, ADMIN", "FT-07", "BF-11"),
    ("UC-16", "Manage Quality Forms", "ADMIN", "FT-08", "BF-05"),
    ("UC-17", "Manage Quality Form Criteria", "ADMIN", "FT-08/FT-09", "BF-05"),
    ("UC-18", "Manage Quality Inspection History", "USER, MANAGER, ADMIN", "FT-08", "BF-06"),
    ("UC-19", "Manage Quality Indicator Formulas", "ADMIN", "FT-09", "BF-06"),
    ("UC-20", "Configure Quality Targets", "ADMIN", "FT-10", "BF-06"),
    ("UC-21", "Manage Quality Alerts", "ADMIN", "FT-10", "BF-10"),
    ("UC-22", "Analyze Quality Dashboard & Reports", "USER, MANAGER, ADMIN", "FT-11", "BF-11"),
    ("UC-23", "Manage Competency Tests", "ADMIN / evaluation permissions", "FT-12", "BF-07"),
    ("UC-24", "Manage Question Bank", "QUESTION_AUTHOR, QUESTION_REVIEWER, ADMIN", "FT-12", "BF-07"),
    ("UC-25", "Configure Test Settings", "EXAM_PUBLISHER, ADMIN", "FT-12", "BF-07"),
    ("UC-26", "Execute Competency Test", "USER", "FT-13", "BF-08"),
    ("UC-27", "Track Test Results", "USER, MANAGER, ADMIN", "FT-13", "BF-08"),
    ("UC-28", "Manage Competency Classification Rules", "ADMIN / evaluation permissions", "FT-14", "BF-08"),
    ("UC-29", "Analyze Performance Dashboard & Reports", "USER, MANAGER, ADMIN", "FT-15", "BF-11"),
    ("UC-30", "Manage Notifications", "USER, MANAGER, ADMIN", "FT-16", "BF-10"),
    ("UC-31", "Manage Email Templates", "ADMIN", "FT-16", "BF-10"),
    ("UC-32", "Manage Question Documents", "QUESTION_AUTHOR, QUESTION_REVIEWER, ADMIN", "FT-17", "BF-09"),
    ("UC-33", "Review Generated Question Candidates", "QUESTION_REVIEWER, ADMIN", "FT-17", "BF-09"),
    ("UC-34", "Manage Prompt Templates", "ADMIN / evaluation permissions", "FT-17", "BF-09"),
]

BFS = [
    ("BF-01", "Account & Access Lifecycle", "Administrator provisions an account; the employee completes first-login verification, signs in and receives role/permission-scoped access."),
    ("BF-02", "Master Data & System Setup", "Administrator maintains departments, positions, roles, settings and reference data used by all modules."),
    ("BF-03", "Training Record & Evidence Lifecycle", "Employee records CME activity, attaches evidence and submits; authorized staff review the record and audit trail."),
    ("BF-04", "Training Compliance Monitoring", "The system totals valid submitted hours in the configured window, identifies gaps and presents scoped compliance views."),
    ("BF-05", "Quality Form Design & Assignment", "Administrator designs and versions a checklist, configures scoring and assigns it to employees or departments."),
    ("BF-06", "Quality Inspection & Scoring", "Assignee completes the checklist; the system validates, scores, compares targets and retains versioned history."),
    ("BF-07", "Competency Test Authoring & Publishing", "Authorized staff curate questions, configure an exam, generate a paper and assign it to employees."),
    ("BF-08", "Competency Test Execution & Classification", "Employee completes an attempt; the system grades on a 0–10 scale and updates competency classification."),
    ("BF-09", "AI Question Generation & Review", "Author uploads a document; asynchronous generation/deduplication produces candidates that a reviewer accepts or rejects."),
    ("BF-10", "Notification & Email Dispatch", "Domain events create notifications and queued emails; retryable delivery records keep failures observable."),
    ("BF-11", "Analytics & Reporting", "Authorized users filter training, quality and competency dashboards and export allowed results."),
    ("BF-12", "Legacy Data Import & Audit", "Administrator previews structured files, resolves validation warnings, applies confirmed rows and reviews the import log."),
]

GBRS = [
    ("GBR-01", "All protected operations require a valid JWT; authorization is enforced at routes and service/controller methods."),
    ("GBR-02", "Roles are ADMIN, MANAGER and USER; evaluation permissions form an additional authorization dimension."),
    ("GBR-03", "Access tokens expire after 15 minutes and refresh tokens after 7 days; refresh is rotated."),
    ("GBR-04", "Training records are editable only in DRAFT."),
    ("GBR-05", "TrainingRecord transitions are DRAFT→SUBMITTED/CANCELLED and admin-only SUBMITTED→DRAFT/CANCELLED; CANCELLED is terminal."),
    ("GBR-06", "Compliance hours count valid SUBMITTED records in the configured window."),
    ("GBR-07", "Evidence upload must satisfy type, size, checksum and storage authorization checks."),
    ("GBR-08", "Form versions used by submissions remain auditable; scoring changes run through recalculation jobs."),
    ("GBR-09", "Exam scores and passing thresholds use the 0–10 scale end to end."),
    ("GBR-10", "Managers see only their allowed department scope; users see personal data unless explicitly permitted."),
    ("GBR-11", "AFTER_COMMIT write handlers use a new transaction to avoid writes joining a completed transaction."),
    ("GBR-12", "Asynchronous generation/paraphrase/recalculation jobs expose state and must be idempotent for retries."),
    ("GBR-13", "Error responses use a stable error code, correlation ID and Vietnamese user-facing message."),
    ("GBR-14", "Secrets remain outside source control and must never be returned to the browser or logs."),
    ("GBR-15", "UAT sign-off is required before the delivery can be labelled formally accepted."),
]

NFRS = [
    ("NFR-P01", "Performance", "Common API requests should complete within 2 seconds under the agreed normal-load profile.", "k6 load script + percentile report", "Pending controlled environment run"),
    ("NFR-P02", "Performance", "Scoring and dashboard queries remain usable under concurrent activity without incorrect totals.", "k6 scoring script + reconciliation", "Pending controlled environment run"),
    ("NFR-S01", "Security", "Protected APIs reject missing, expired and unauthorized tokens.", "L2/L3 security tests", "Implemented; regression evidence maintained"),
    ("NFR-S02", "Security", "Passwords and secrets are never stored or logged in plaintext.", "configuration/code review", "Implemented by Spring Security and env configuration"),
    ("NFR-S03", "Security", "Evidence and document objects are served only through authorized application paths.", "integration/security tests", "Implemented"),
    ("NFR-R01", "Reliability", "Queued email and scheduled cleanup work can be retried without duplicate business effects.", "job integration tests", "Implemented with retry/idempotency controls"),
    ("NFR-R02", "Reliability", "Concurrent updates return conflict rather than silently overwriting protected entities.", "optimistic-lock tests", "Implemented on critical versioned entities"),
    ("NFR-U01", "Usability", "All end-user UI text and validation feedback is in Vietnamese.", "UI review", "Implemented for active product UI"),
    ("NFR-U02", "Usability", "Core staff, manager and admin flows support desktop and mobile-width layouts.", "Playwright responsive suite", "Automation prepared; full environment run pending"),
    ("NFR-M01", "Maintainability", "Backend follows controller/service/repository separation and frontend follows app/features/shared.", "architecture review", "Observed in repository"),
    ("NFR-O01", "Observability", "Errors carry correlation IDs and material business actions are auditable.", "L3 response/audit checks", "Implemented"),
    ("NFR-C01", "Compatibility", "Frontend supports current evergreen browsers; backend deploys on Java 17+ runtime.", "build/browser matrix", "Chrome/Edge baseline; broader matrix pending"),
]

JOBS = [
    ("JOB-01", "Compliance alert scan", "Daily 07:00 Asia/Bangkok", "NotificationAlertScheduler", "Finds below-target users and publishes configured notifications."),
    ("JOB-02", "Evidence object deletion retry", "Every 600,000 ms", "EvidenceObjectDeletionService", "Retries storage deletion for soft-deleted evidence and records completion."),
    ("JOB-03", "Form scoring recalculation dispatcher", "Every 30,000 ms", "FormScoringRecalculationDispatcher", "Claims pending recalculation jobs and dispatches workers."),
    ("JOB-04", "Email queue consumer", "RabbitMQ event-driven", "EmailConsumer", "Consumes email messages with manual acknowledgement and retry/dead-letter behavior."),
    ("JOB-05", "Form scoring worker", "AFTER_COMMIT async", "FormScoringRecalculationWorker", "Recalculates affected submissions outside the caller transaction."),
    ("JOB-06", "Document question generation", "AFTER_COMMIT async", "DocumentQuestionJobWorker", "Chunks documents, calls the configured generator, deduplicates and stores candidates."),
    ("JOB-07", "Question paraphrase", "AFTER_COMMIT async", "ParaphraseJobWorker", "Generates paraphrase candidates using configured runtime and persists job state."),
    ("JOB-08", "Embedding startup backfill/cache warmup", "ApplicationReadyEvent", "QuestionEmbeddingStartupBackfill / EmbeddingCacheWarmup", "Backfills missing embeddings and warms semantic-search cache when enabled."),
]


def routes():
    text = (ROOT / "carehub-frontend/src/app/router.jsx").read_text(encoding="utf-8")
    found = re.findall(r'<Route\s+path="([^"]+)"', text)
    found += re.findall(r'path=\{AUTH_ROUTES\.([A-Za-z0-9_]+)\}', text)
    seen, result = set(), []
    for value in found:
        if value not in seen and value != "*":
            seen.add(value); result.append(value)
    return result


def controller_inventory():
    rows = []
    for path in sorted((ROOT / "carehub-backend/src/main/java").rglob("*Controller.java")):
        text = path.read_text(encoding="utf-8")
        if "@RestController" not in text:
            continue
        base = re.search(r'@RequestMapping\((?:value\s*=\s*)?"([^"]+)"', text)
        endpoints = len(re.findall(r'@(GetMapping|PostMapping|PutMapping|PatchMapping|DeleteMapping)', text))
        rows.append((path.stem, base.group(1) if base else "method-level mappings", endpoints, str(path.relative_to(ROOT)).replace("\\", "/")))
    return rows


def entity_inventory():
    rows = []
    for path in sorted((ROOT / "carehub-backend/src/main/java").rglob("*.java")):
        text = path.read_text(encoding="utf-8")
        if "@Entity" not in text:
            continue
        table = re.search(r'@Table\s*\(\s*name\s*=\s*"([^"]+)"', text)
        rows.append((path.stem, table.group(1) if table else "derived by JPA", str(path.relative_to(ROOT)).replace("\\", "/")))
    return rows


ROUTES = routes()
CONTROLLERS = controller_inventory()
ENTITIES = entity_inventory()


def sha256(path):
    return hashlib.sha256(path.read_bytes()).hexdigest()


def distill(template):
    doc = Document(template)
    target = ARTIFACTS / template.stem
    target.mkdir(parents=True, exist_ok=True)
    sections = []
    for s in doc.sections:
        sections.append(f"- {s.page_width/914400:.2f}×{s.page_height/914400:.2f} in; margins L/R/T/B {s.left_margin/914400:.2f}/{s.right_margin/914400:.2f}/{s.top_margin/914400:.2f}/{s.bottom_margin/914400:.2f} in")
    headings = [p.text.strip() for p in doc.paragraphs if p.text.strip() and p.style.name.startswith("Heading")]
    package_parts = []
    import zipfile
    with zipfile.ZipFile(template) as z:
        for info in z.infolist():
            package_parts.append(f"- {info.filename}: {info.file_size} bytes")
    content = f"""# Retained template contract

Reference: `{template.resolve()}`  
SHA-256: `{sha256(template)}`  
Sections: {len(doc.sections)}; body paragraphs: {len(doc.paragraphs)}; tables: {len(doc.tables)}

## Page system
{chr(10).join(sections)}

## Typography and components
- Preserve source styles, numbering definitions, theme, headers, footers, section geometry and relationships.
- Reuse Heading 1–4, Normal and source table styles; do not shrink body text to fit.
- Tables repeat their header row, use source-derived blue header fills, fixed readable widths and page-safe wrapping.

## Content flow
{chr(10).join('- ' + h for h in headings)}

## Slot map
- The template instructional/example body is replaceable because this task creates the complete CareHub report.
- Theme, styles, numbering, headers, footers and section properties are preserve-only.
- Title, metadata, change log and section bodies are rewritten with CareHub facts.

## Package preservation
{chr(10).join(package_parts)}

## Fidelity gates
- Reference SHA must remain unchanged.
- Output must keep section geometry and recognizable source typography.
- Render every output page; reject clipping, overlap, broken tables, placeholder text or template examples.
"""
    (target / "artifact.md").write_text(content, encoding="utf-8")


def clear_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    try: table.style = "Table Grid"
    except Exception: pass
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        set_cell_shading(cell, "1F4E78")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True; run.font.color.rgb = RGBColor(255,255,255); run.font.size = Pt(8)
    repeat_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs: run.font.size = Pt(8)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths): row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text); run.bold = bold
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.6)
        para.paragraph_format.first_line_indent = Cm(-0.4)
        para.add_run("• ").bold = True
        para.add_run(item)


def add_figure(doc, image_name, caption, alt_text, width_cm=15.8):
    image_path = DIAGRAMS / image_name
    if not image_path.exists():
        raise FileNotFoundError(f"Missing report diagram: {image_path}")
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.keep_with_next = True
    shape = para.add_run().add_picture(str(image_path), width=Cm(width_cm))
    shape._inline.docPr.set("descr", alt_text)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(9)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(89, 89, 89)
    return shape


def cover(doc, title, report):
    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(54); para.paragraph_format.space_after = Pt(18)
    run = para.add_run(title); run.bold = True; run.font.size = Pt(24); run.font.color.rgb = RGBColor(31,78,120)
    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(f"{PROJECT} — {report}"); r.font.size = Pt(13); r.bold = True
    add_table(doc, ["Document field", "Value"], [
        ("Version", VERSION), ("Date", DATE), ("Authors", AUTHORS),
        ("Reviewer", "Nguyễn Văn An — Supervisor"), ("Status", "Draft for supervisor review; UAT sign-off pending"),
        ("Evidence baseline", "Repository branch ManhTuan, commit 627b9448; templates 26FA v0.9"),
    ], [4.2, 12.0])
    doc.add_page_break()
    doc.add_heading("Change Log", level=1)
    add_table(doc, ["Version", "Date", "Author", "Summary"], [(VERSION, DATE, "CareHub team", "Rebuilt from the 26FA v0.9 template and reconciled with current source code, prior reports, tracking and test evidence.")], [2,2.5,4,8])
    p(doc, "Document control note: source code is authoritative for implemented behavior. UAT and production acceptance remain pending until signed by the Product Owner.")


def new_doc(template_name, output_name, title, report):
    template = TEMPLATES / template_name
    distill(template)
    doc = Document(template)
    clear_body(doc)
    doc.core_properties.title = title
    doc.core_properties.subject = PROJECT
    doc.core_properties.author = AUTHORS
    doc.core_properties.comments = "Generated from the retained 26FA v0.9 template; evidence-grounded against the CareHub repository."
    cover(doc, title, report)
    return doc, OUT / output_name


def save(doc, path):
    path.parent.mkdir(parents=True, exist_ok=True)
    for section in doc.sections:
        for story in (section.header, section.footer):
            for node in story._element.iter():
                if node.tag == qn("w:t") and node.text:
                    node.text = node.text.replace("{{Project Code}}", "CareHub")
    body = doc._element.body
    for child in reversed(list(body)):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:p") and not "".join(child.itertext()).strip():
            body.remove(child)
            continue
        break
    settings = doc.settings._element
    if settings.find(qn("w:updateFields")) is None:
        node = OxmlElement("w:updateFields"); node.set(qn("w:val"), "true"); settings.append(node)
    doc.save(path)


def build_r1():
    d, out = new_doc("Report 1_VS_Template.docx", "Report 1_VS_CareHub_v0.9.docx", "Vision & Scope Document", "Report 1")
    d.add_heading("1. Product Background", 1)
    p(d, "Viet Duc Hospital needs one auditable workspace for continuing medical education, care-quality checklists and competency assessment. The previous combination of spreadsheets, evidence folders and disconnected forms makes it difficult to confirm training compliance, compare quality results and identify capability gaps at employee or department level.")
    add_table(d,["Pain area","Current situation","Operational consequence"],[
        ("Training compliance","Records and evidence are distributed across files and manual workflows.","Managers cannot reliably reconcile hours, validity and compliance."),
        ("Quality inspection","Forms and scoring rules vary and results are hard to compare.","Benchmark breaches and corrective priorities are detected late."),
        ("Competency assessment","Questions, attempts and classification are not governed in one lifecycle.","Retake history and capability gaps lack a single audit trail."),
        ("Reporting","Metrics require manual consolidation.","Management decisions are delayed and totals may diverge."),
    ],[3,6,7])
    d.add_heading("2. Existing Solutions",1)
    add_table(d,["Solution","Strength","Gap for CareHub"],[
        ("SafetyCulture","Mature inspection/checklist and evidence workflows.","General-purpose model; no integrated CME and competency lifecycle."),
        ("HealthStream","Healthcare learning, compliance and workforce development.","Does not cover the hospital-specific quality-form/scoring workflow required here."),
        ("RLDatix","Healthcare safety, audit and governance capabilities.","Not primarily a training-hour and employee-testing platform."),
        ("Spreadsheets and shared folders","Low entry cost and familiar tools.","Weak access control, versioning, traceability, automation and cross-module reporting."),
    ],[3.2,6,7])
    d.add_heading("3. Proposed Solution",1)
    p(d, "CareHub is a Vietnamese-language web application for staff, department managers, administrators and delegated evaluation roles. It joins three operational pillars—training, quality evaluation and competency assessment—under common identity, reference data, notifications, auditability and analytics. The application is implemented as a React SPA backed by a Spring Boot REST API, PostgreSQL, RabbitMQ and object storage.")
    add_figure(d,"01_system_context.png","Figure 1. CareHub system context and external dependencies.","CareHub system context showing hospital users, the CareHub system boundary, data stores, messaging, file storage and AI services.")
    d.add_heading("4. Project Scope & Limitations",1); d.add_heading("4.1 Major Features",2)
    add_table(d,["ID","Feature","Description","Gap"],[(x[0],x[1],x[6],x[5]) for x in FEATURES],[2,4.5,8,2.5])
    d.add_heading("4.2 Limitations & Exclusions",2)
    add_table(d,["ID","Excluded item","Reason"],[
        ("LI-01","Direct integration with hospital core systems","No approved production API/database access is available for the project baseline."),
        ("LI-02","Native iOS/Android application","The responsive web application is the delivery target."),
        ("LI-03","SMS/Zalo gateway","Current communication uses in-app notification and email."),
        ("LI-04","Fully automated clinical decision-making","AI only proposes questions; human review remains mandatory."),
        ("LI-05","Formal production acceptance","Requires executed UAT, Product Owner signature and customer IT deployment."),
    ],[2,6,9])
    d.add_heading("5. Expected Contributions & Next Steps",1)
    d.add_heading("5.1 Business & Operational Impact",2); p(d,"CareHub replaces repeated manual reconciliation with traceable records, state transitions and shared dashboards. Success is measured through complete requirement coverage, stable test evidence, fewer manual consolidations and faster identification of below-target staff or departments.")
    d.add_heading("5.2 User Experience Impact",2); p(d,"Staff can see personal training, checklist, exam and competency information in one interface. Managers receive department-scoped views; administrators manage shared configuration and audit-sensitive operations without exchanging uncontrolled spreadsheets.")
    d.add_heading("5.3 Planned Next Steps After Project Completion",2); bullets(d,["Execute and sign the 12 UAT business-flow scripts.","Run controlled k6 performance tests and record p95 results.","Deploy to staging with production-like RabbitMQ, PostgreSQL and R2 configuration.","Complete security review and secret rotation before go-live.","Collect pilot feedback and prioritize hospital integration work for the next version."])
    save(d,out)


def build_r2():
    d,out=new_doc("Report 2.0_ProjectPlan_Template.docx","Report 2.0_ProjectPlan_CareHub_v0.9.docx","Project Management Plan","Report 2")
    d.add_heading("1. Project Overview",1); d.add_heading("1.1 System Description",2); p(d,"CareHub supports hospital staff development and quality governance through connected training, checklist, examination, notification and reporting workflows.")
    d.add_heading("1.2 Technical Scope Summary",2)
    add_table(d,["Metric","Baseline"],[("Business flows",f"{len(BFS)} (BF-01–BF-12)"),("Actors/roles","ADMIN, MANAGER, USER plus evaluation permissions"),("Use cases",f"{len(UCS)} (UC-01–UC-34)"),("Feature groups",f"{len(FEATURES)} (FT-01–FT-17)"),("Frontend routes",str(len(ROUTES))),("REST controllers/endpoints",f"{len(CONTROLLERS)} / 302 mapping annotations"),("Persistent entities",str(len(ENTITIES))),("Background jobs",f"{len(JOBS)}")],[5,11])
    d.add_heading("2. Development Team",1); d.add_heading("2.1 Team Structure",2)
    add_table(d,["Member","Primary role","Responsibilities"],[
        ("Đỗ Mạnh Tuấn","Tech Lead / PM","Architecture, backend, integration, deployment, document baseline."),("Dương Ngọc Hiếu","Backend / Data","Services, persistence, import/export, integration tests."),("Ngô Trung Kiên","BA / QA","Requirements, traceability, scenario and acceptance design."),("Phí Hoàng Nam","Frontend / QA","Admin/evaluation UI, usability and regression tests."),("Phạm Hải Nam","Frontend / QA","Staff/manager UI, responsive behavior and UAT support."),
    ],[4,4,9])
    d.add_heading("2.2 RACI by Deliverable Type",2)
    add_table(d,["Deliverable","Tuấn","Hiếu","Kiên","H. Nam","Hải Nam"],[("VS/PRD/UCS","A/C","C","R","C","C"),("TDS/Deployment","A/R","R","C","I","I"),("FDS/User Manual","A/C","C","C","R","R"),("L1/L2","A","R","C","R","R"),("L3/UAT","A","C","R","R","R"),("Final report","A/R","C","R","C","C")],[4,2.4,2.4,2.4,2.4,2.4])
    d.add_heading("2.3 Capacity Planning",2); add_table(d,["Item","Value"],[ ("Nominal team","5 members"),("Nominal work week","40 hours/person"),("Planning duration","14 weeks"),("Gross capacity","2,800 hours"),("20% coordination/review buffer","560 hours"),("Effective planning capacity","2,240 hours"),("Legacy WBS estimate","2,612 hours — retained as historical plan, not claimed actual")],[6,10])
    d.add_heading("3. Deliverable & Scope",1); d.add_heading("3.1 Timeline Overview",2)
    add_table(d,["Phase","Weeks","Outcome"],[ ("Plan / analyze / design","1–5","VS, project plan, PRD/UCS baseline, architecture and test plan"),("Iteration 1","6–7","Access, administration and training foundation"),("Iteration 2","8–10","Quality workflows, training compliance and manager views"),("Iteration 3","11–12","Evaluation, AI jobs, analytics and notifications"),("Verification / handover","13–14","L3/E2E/UAT package, deployment/user guides and final report")],[4,3,10])
    add_figure(d,"02_project_lifecycle.png","Figure 1. Project delivery lifecycle and release gates.","Five-stage delivery lifecycle from discovery and specification through design, implementation, verification, acceptance and handover.")
    d.add_heading("3.2 Project Deliverables",2)
    add_table(d,["Deliverable","Owner","Status"],[(f"Report {i}","Team","Rebuilt in 26FA v0.9 package") for i in ["1 VS","2.0 Plan","2.1 Tracking","3.0 PRD","3.1 UCS","4.0 TDS","4.1 FDS","5.0 Test","5.1–5.4 Workbooks","6.1 Deployment","6.2 User Manual","7 Final"]],[7,4,6])
    d.add_heading("3.3 Scope and Workload",2)
    for title, modules in [("a. Iteration 1",FEATURES[:5]),("b. Iteration 2",FEATURES[5:11]),("c. Iteration 3",FEATURES[11:])]:
        d.add_heading(title,3); add_table(d,["FT-ID","Feature","Priority"],[(x[0],x[1],x[2]) for x in modules],[2.5,11,3])
    d.add_heading("4. Risks and Mitigations",1)
    add_table(d,["ID","Risk","P/I","Mitigation"],[
        ("R-01","Hospital integration access unavailable","M/H","Use controlled import; isolate adapters for later API integration."),("R-02","Production-like messaging/storage differs from local","M/H","Validate staging with RabbitMQ/R2 and smoke/retry checks."),("R-03","Large SPA bundle affects initial load","H/M","Code-split high-cost evaluation routes; track bundle size."),("R-04","UAT owner unavailable or scripts incomplete","M/H","Book UAT window early; keep Critical/High scripts first."),("R-05","Requirements/docs diverge from code","H/H","Use immutable IDs and regenerate tracker/reports from one baseline."),("R-06","AI models/API unavailable","M/M","Expose model health, circuit-breaker fallback and manual authoring path."),
    ],[2,7,2,7])
    d.add_heading("5. Communication Plan",1); add_table(d,["Cadence","Participants","Purpose","Evidence"],[("Daily async","Team","Progress/blockers","Team channel"),("Weekly review","Team + supervisor","Scope, risk, decisions","Meeting notes"),("Per pull request","Developers/reviewer","Code and test review","Git history"),("End of iteration","Team + supervisor","Demo and acceptance evidence","Test reports"),("UAT session","Product Owner + QA","Business acceptance","Signed UAT workbook")],[4,5,6,4])
    d.add_heading("6. Configuration Management",1); d.add_heading("6.1 Document Management",2); p(d,"Official documents use the 26FA v0.9 filenames and immutable requirement IDs. Templates and docs-old remain read-only; the completed package is stored under docs/26FA_v0.9_CareHub.")
    d.add_heading("6.2 Source Code Management",2); p(d,"Git branches and reviewed pull requests are the code record. Release claims reference a commit hash; secrets and generated runtime artifacts are excluded.")
    d.add_heading("6.3 Tools & Infrastructures",2); add_table(d,["Area","Technology"],[("Backend","Java 17; Spring Boot 4.0.6; Maven"),("Frontend","React 19.2.6; Vite 8; React Router 7; Axios; Recharts"),("Data","PostgreSQL 17; JPA/Hibernate"),("Async/cache","RabbitMQ 3; Redis 7"),("Storage/AI","Cloudflare R2; DeepSeek API; multilingual-e5-small; VietQuill"),("Testing","JUnit 5, Spring Boot Test, Vitest, Testing Library, Playwright, k6, JaCoCo"),("Deployment","Docker Compose, reverse proxy/TLS, environment variables")],[5,12])
    save(d,out)


def build_prd():
    d,out=new_doc("Report 3.0_PRD_Template.docx","Report 3.0_PRD_CareHub_v0.9.docx","Product Requirements Document","Report 3.0")
    d.add_heading("1. Business Flows",1); d.add_heading("1.1 Overview",2); p(d,"The flows below are the stable business-level spine used by UCS, FDS, TestDoc, UAT and Project Tracking.")
    add_figure(d,"03_feature_map.png","Figure 1. Product capability map and shared governance foundation.","CareHub capability map showing training, quality evaluation and competency pillars supported by shared authentication, reference data, notifications, audit and analytics.")
    add_table(d,["BF-ID","Flow","Business outcome"],BFS,[2.2,5,10])
    for bf,name,desc in BFS:
        d.add_heading(f"{bf}: {name}",2); p(d,desc); add_table(d,["Stage","Actor / system responsibility"],[ ("Trigger","Authorized actor starts the workflow from an allowed screen."),("Validate","UI and API validate identity, scope, state and input."),("Process","Domain service applies the transaction and audit rules."),("Complete","User sees the persisted end state; asynchronous follow-up remains observable."),("Exception","Failure returns a Vietnamese message and correlation ID without partial business state.")],[3,14])
    d.add_heading("2. Conceptual Data Model",1); d.add_heading("2.1 Entity Relationship Overview",2); p(d,"User/Department/Role provide identity and scope. TrainingRecord/Evidence/ActivityType support CME. Form/FormVersion/Assignment/Submission support quality evaluation. Question/ExamPaper/Attempt and competency entities support assessment. Notification, ImportLog and AuditLog provide cross-cutting traceability.")
    d.add_heading("2.2 Entity List",2); add_table(d,["Entity","Physical table","Source"],ENTITIES,[5,5,8])
    d.add_heading("2.3 Data Business Rules",2); add_table(d,["Rule","Requirement"],GBRS[:12],[2.5,14])
    d.add_heading("3. User Requirements",1); d.add_heading("3.1 Actor Definitions",2); add_table(d,["Actor","Need"],[ ("Staff (USER)","Manage personal training, checklists, exams, notifications and competency views."),("Department Manager","Monitor and act within department scope."),("Administrator","Manage accounts, reference data, settings and hospital-wide configuration."),("Evaluation Author/Reviewer/Publisher","Perform delegated evaluation work without requiring full administrator access."),("Product Owner / Supervisor","Review business outcomes and sign UAT acceptance.")],[5,12])
    d.add_heading("3.2 Permission Matrix",2); add_table(d,["Capability","USER","MANAGER","ADMIN","Evaluation permissions"],[ ("Personal records","Own","Department + own","All","As permitted"),("Reference/settings","Read limited","Read limited","Manage","No"),("Quality/evaluation","Assigned","Department","All","Delegated scope"),("Reports","Personal","Department","Hospital","Delegated evaluation"),("Account administration","No","No","Yes","No")],[6,2.5,3,3,4])
    d.add_heading("3.3 Use Case Index",2); add_table(d,["UC-ID","Use case","Actors","FT-ID","BF-ID"],UCS,[2,5.5,5.5,3,3])
    d.add_heading("4. Functional Requirements",1); d.add_heading("4.1 Modules & Features",2); add_table(d,["FT-ID","Feature","Priority","Module","UC","BF","Requirement"],FEATURES,[2,4,2,3,3,3,8])
    d.add_heading("4.2 Screen Inventory",2); add_table(d,["SCR-ID","Route","Delivery"],[(f"SCR-{i:03d}",r,"Implemented route declaration") for i,r in enumerate(ROUTES,1)],[2.5,11,5])
    d.add_heading("4.3 External API Inventory",2); add_table(d,["API group","Base mapping","Methods","Source"],CONTROLLERS,[5,6,2,7])
    d.add_heading("4.4 Background Job Inventory",2); add_table(d,["JOB-ID","Name","Trigger","Implementation","Outcome"],JOBS,[2,4,4,5,7])
    d.add_heading("5. Global Business Rules",1); add_table(d,["GBR-ID","Rule"],GBRS,[2.5,15])
    d.add_heading("6. Non-Functional Requirements & Constraints",1); add_table(d,["NFR-ID","Category","Requirement","Verification","Status"],NFRS,[2.5,3,7,5,5])
    save(d,out)


def build_ucs():
    d,out=new_doc("Report 3.1_UCS_Template.docx","Report 3.1_UCS_CareHub_v0.9.docx","Use Case Specification","Report 3.1")
    d.add_heading("Glossary",1); add_table(d,["Term","Meaning"],[("CME","Continuing Medical Education"),("R2","Cloudflare object storage used for evidence/documents"),("Evaluation permission","Fine-grained authority in addition to ADMIN/MANAGER/USER"),("Main Flow","Successful actor/system sequence"),("Alternative Flow","Validation, authorization or dependency branch")],[5,12])
    d.add_heading("Conventions & Notation",1); d.add_heading("Entity State Machines",2); add_table(d,["Entity","States / rule"],[ ("TrainingRecord","DRAFT → SUBMITTED or CANCELLED; admin-only SUBMITTED → DRAFT/CANCELLED; CANCELLED terminal"),("ExamAttempt","Assigned/available → in progress → submitted/graded; score is 0–10"),("Async job","Pending → Running → Completed/Failed; retry must not duplicate accepted output")],[5,12])
    d.add_heading("Step Notation",2); p(d,"A = actor action; S = synchronous system response; J = asynchronous job/event. Verification criteria use Given/When/Then and reference the L1–L4 workbooks.")
    d.add_heading("Use-Case Overview Diagrams",1)
    add_figure(d,"04_use_case_landscape.png","Figure 1. Actor-to-use-case landscape.","Overview mapping CareHub actors to access, training, quality, competency, analytics and operational use-case groups.")
    add_figure(d,"05_business_flows.png","Figure 2. Business-flow coverage grouped by domain.","Twelve CareHub business flows grouped into access, training, quality, competency and operational domains.")
    for uc,name,actors,ft,bf in UCS:
        d.add_heading(f"{uc}: {name}",1)
        add_table(d,["Field","Value"],[("Related feature",ft),("Business flow",bf),("Primary actors",actors),("Trigger",f"Authorized actor selects the {name.lower()} capability."),("Preconditions","Account is active; required role/permission and reference data exist."),("Success end state",f"The requested {name.lower()} change is persisted or displayed and is auditable."),("Failure guarantee","No unauthorized or partial state is committed; user receives actionable Vietnamese feedback.")],[4,13])
        d.add_heading("Main Flow",2); add_table(d,["Step","Actor/System","Action"],[(1,"Actor","Opens the permitted screen or operation."),(2,"System","Loads scoped data and current entity version."),(3,"Actor","Enters or confirms the required data."),(4,"System","Validates input, authorization, state and business rules."),(5,"System","Commits the transaction and audit/event record."),(6,"System","Returns the resulting view/state; any background work exposes a job or notification state.")],[2,4,11])
        d.add_heading("Alternative Flows",2); add_table(d,["ID","Condition","Expected behavior"],[(f"AF-{uc[3:]}-01","Authentication/authorization fails","Return 401/403 or guarded navigation; disclose no restricted data."),(f"AF-{uc[3:]}-02","Input or state rule fails","Return validation/conflict response; keep prior state."),(f"AF-{uc[3:]}-03","External storage, email or AI dependency fails","Record retryable failure or use documented fallback; preserve transaction integrity.")],[3,6,8])
        d.add_heading("Business Rules",2); p(d,f"Applies {ft}, {bf} and global rules GBR-01, GBR-02, GBR-10, GBR-13; domain-specific rules are traced in PRD §5.")
        d.add_heading("Request Fields / Data",2); add_table(d,["Field group","Rule"],[("Identifiers","Positive existing IDs; scope checked server-side."),("Text","Trimmed, length-limited and validated; UI validation does not replace server validation."),("Dates/numbers","Typed values with domain boundaries; score remains 0–10 where applicable."),("Version","Required for optimistic-lock protected updates."),("Files","Allowed MIME/extension/size; checksum and object authorization applied.")],[5,12])
        d.add_heading("Allowed Roles",2); p(d,actors)
        d.add_heading("Verification Criteria",2); add_table(d,["Criterion","Evidence"],[("Happy path reaches the stated end state",f"L2 integration + L3 flow mapped to {uc}"),("Each listed alternative prevents invalid state","Negative/security cases in L1/L2/L3"),("Business outcome is understandable to stakeholders",f"UAT script mapped to {bf}")],[9,8])
    save(d,out)


def build_tds():
    d,out=new_doc("Report 4.0_TDS_Template.docx","Report 4.0_TDS_CareHub_v0.9.docx","Technical Design Specification","Report 4.0")
    d.add_heading("1. Tech Stack",1); d.add_heading("1.1 Technology Choices",2); add_table(d,["Layer","Technology","Version / note"],[("Backend","Java / Spring Boot","Java 17 target; Spring Boot 4.0.6"),("Frontend","React / Vite","React 19.2.6; Vite 8.x; JavaScript/JSX"),("Routing/API","React Router / Axios","7.18.1 / 1.17.x"),("Database","PostgreSQL","17; Hibernate ddl-auto update"),("Messaging/cache","RabbitMQ / Redis","3 / 7"),("Storage","Cloudflare R2","S3-compatible evidence/document objects"),("AI","DeepSeek + E5 + VietQuill","API generation; ONNX embedding; local T5 paraphrase"),("Testing","JUnit/Vitest/Playwright/k6","Unit through system/browser/performance")],[4,6,8])
    d.add_heading("1.2 External Integrations",2); add_table(d,["Integration","Purpose","Failure handling"],[("SMTP via RabbitMQ","Email delivery","Queued retry and observable failure"),("Cloudflare R2","Evidence/document storage","Compensating delete + retry sweep"),("DeepSeek API","Question generation","Timeout/retry/circuit breaker and fallback"),("Local model files","Embedding/paraphrase","Health state and manual-authoring fallback")],[5,6,7])
    d.add_heading("1.3 Decision Rationale",2); p(d,"A modular monolith keeps transaction and deployment complexity appropriate for a five-person team while separating domains by package. The SPA/API boundary supports the route-rich role experience; asynchronous workers isolate slow email, scoring and AI tasks.")
    d.add_heading("2. Architecture Overview",1); d.add_heading("2.1 High-Level Architecture",2); p(d,"Browser → React SPA → Axios/JWT → Spring Security → REST Controller → Service → Repository → PostgreSQL. Services publish AFTER_COMMIT events to async workers or RabbitMQ; storage and AI adapters isolate external dependencies.")
    add_figure(d,"06_architecture.png","Figure 1. High-level technical architecture.","Layered CareHub architecture from browser and React SPA through Spring API and domain services to PostgreSQL, Redis, RabbitMQ, R2 and AI adapters.")
    d.add_heading("2.2 Layer Responsibilities",2); add_table(d,["Layer","Responsibility"],[("UI","Routing, forms, role guards, client feedback"),("Controller","HTTP contract, request validation, method authorization"),("Service","Transactions, business rules, state machines, orchestration"),("Repository","Scoped persistence queries and optimistic locking"),("Async/integration","Email, recalculation, AI jobs, storage cleanup"),("Common","Response envelope, errors, correlation, shared utilities")],[5,12])
    d.add_heading("2.3 Package Convention",2); p(d,"Backend packages: auth, user, training, form, questiongeneration, notification, imports, dashboard, systemsettings, config, common and exception. Frontend packages: app → features → shared; pages call feature API modules rather than hard-coded URLs.")
    d.add_heading("2.4 Request Lifecycle",2); bullets(d,["Request interceptor attaches the access token.","Spring Security decodes JWT roles/permissions and method rules authorize the operation.","Controller validates DTO and delegates to a transactional service.","Service enforces scope/state, writes domain and audit state, then publishes follow-up events.","GlobalExceptionHandler maps failure to error_code/message/correlation_id/details.","Client shows Vietnamese feedback; 401 refresh is deduplicated and retried once."])
    d.add_heading("3. Data Model",1); d.add_heading("3.1 Entity Definitions",2); add_table(d,["Entity","Table","Source"],ENTITIES,[5,5,8])
    d.add_heading("3.2 Entity Relationships",2); p(d,"Users belong to departments and roles; training records own evidence/change logs; form templates own versions/sections/items and feed assignments/submissions; question documents create generation jobs/candidates; questions feed sets, papers, assignments and attempts; notifications and audit records reference domain actors/events.")
    add_figure(d,"07_domain_map.png","Figure 2. Domain data map and cross-context relationships.","Domain data map connecting identity and reference data with training, quality forms, competency assessment and operational records.")
    d.add_heading("3.3 Enum / Lookup Values",2); add_table(d,["Area","Values"],[("Role","ADMIN, MANAGER, USER"),("TrainingRecordStatus","DRAFT, SUBMITTED, CANCELLED"),("Compliance","COMPLIANT, NON_COMPLIANT"),("Async jobs","PENDING, RUNNING, COMPLETED, FAILED variants by job type"),("Priorities","Must, Should, Could; Critical, High, Medium, Low")],[5,12])
    d.add_heading("3.4 Indexing Strategy",2); p(d,"Use primary/foreign keys, unique business identifiers, status/time filters and department/employee composite paths. High-volume dashboard and semantic-search queries require bounded paging and measured indexes before production.")
    d.add_heading("3.5 Migration Strategy",2); p(d,"Current production configuration uses Hibernate ddl-auto=update; SQL fixes exist for targeted constraints. Before hospital production, baseline a reviewed schema and adopt versioned forward migrations with backup/restore rehearsal.")
    d.add_heading("3.6 Seed Data",2); p(d,"Configurable admin account, hospital question bank and nursing professional fields are seeded when app.seed.enabled is true. Production credentials must be provided through environment variables.")
    d.add_heading("4. Security Design",1); add_table(d,["Control","Implementation"],[("Authentication","Stateless OAuth2 Resource Server JWT, HS256, 15-minute access and 7-day refresh"),("Authorization","Route guards plus @PreAuthorize and evaluation permission checks"),("Passwords/secrets","Encoded passwords; secrets from non-committed environment configuration"),("Object access","Evidence/document access through authorized endpoints; internal object keys not exposed"),("Concurrency","@Version on critical mutable entities; conflicts map to 409"),("CORS","Local patterns plus approved production origin"),("Errors","Stable codes and correlation IDs; no stack trace to client")],[5,12])
    d.add_heading("5. UI/UX Conventions",1); bullets(d,["Vietnamese user-facing text and actionable error messages.","Role-aware navigation with ProtectedRoute guards.","Feature-scoped CSS and shared table/UI styles.","Responsive pages for desktop and mobile widths.","Loading, empty and retry states are explicit; destructive actions require confirmation.","Keyboard-visible controls and semantic labels are required for final accessibility review."])
    d.add_heading("6. Configuration & Environment",1); add_table(d,["Group","Key examples"],[("Server/API","SERVER_PORT; app.api-prefix=/api/v1"),("Database","DB_URL, DB_USERNAME, DB_PASSWORD"),("JWT","JWT_SECRET"),("Mail/RabbitMQ","MAIL_*; RABBITMQ_*"),("Storage","R2_ENDPOINT, R2_ACCESS_KEY, R2_SECRET_KEY, BUCKET_NAME"),("AI","GENERATION_*, E5_*, VIETQUILL_*"),("Frontend","VITE_API_BASE_URL")],[5,12])
    d.add_heading("7. File & Storage Conventions",1); p(d,"Evidence and question documents are validated before storage, referenced by opaque object keys, downloaded through authorized application endpoints and removed through soft-delete plus AFTER_COMMIT cleanup/retry. Maximum HTTP multipart file size is 20 MB; feature-specific limits may be stricter.")
    d.add_heading("8. Error Handling & Logging",1); add_table(d,["Code","Meaning"],[ ("REQ_001","Bad request"),("VAL_001","Validation failure"),("AUTH_001","Authentication failure"),("AUTH_002","Forbidden"),("SYS_404","Not found"),("SYS_409","Conflict"),("SYS_503","Dependency unavailable"),("SYS_001","Internal error")],[3,14])
    add_figure(d,"08_request_event.png","Figure 3. Synchronous request and asynchronous event lifecycle.","Request lifecycle through authentication, controller validation, domain transaction, database commit, after-commit worker and external outcome.")
    d.add_heading("9. Performance Constraints",1); add_table(d,["NFR","Design response"],[(x[0],x[2]) for x in NFRS if x[1] in {"Performance","Reliability"}],[3,14]); p(d,"Known watch points: the current production SPA bundle is approximately 1.8 MB minified and Vite reports chunks over 500 kB; evaluation route code-splitting is a release optimization item.")
    save(d,out)


def build_fds():
    d,out=new_doc("Report 4.1_FDS_Template.docx","Report 4.1_FDS_CareHub_v0.9.docx","Feature Design Specification","Report 4.1")
    d.add_heading("1. Document Overview",1); d.add_heading("1.1 Notation & Conventions",2); p(d,"SCR identifies a declared frontend route, API identifies a controller group and JOB identifies scheduled/event-driven background work. Redirect-only routes remain in the inventory because they affect navigation behavior.")
    d.add_heading("1.2 System Scope Summary",2); add_table(d,["Inventory","Count"],[("Actors","3 base roles + delegated evaluation permissions"),("Features",len(FEATURES)),("Use cases",len(UCS)),("Declared unique routes",len(ROUTES)),("REST controllers",len(CONTROLLERS)),("Endpoint mapping annotations",302),("Background jobs",len(JOBS))],[8,8])
    d.add_heading("2. Navigation Map",1); d.add_heading("2.1 Screen Index",2); add_table(d,["SCR-ID","Route","Primary audience"],[(f"SCR-{i:03d}",r,"Admin" if r.startswith("/admin") else "Manager" if r.startswith("/manager") else "Staff/all authenticated" if r.startswith(("/staff","/training")) else "Guest/first-login") for i,r in enumerate(ROUTES,1)],[2.5,10,5])
    d.add_heading("2.2 Navigation Flow by Role",2); add_table(d,["Role","Typical flow"],[("Guest","Login → forgot/reset password or first-login email confirmation"),("Staff","Dashboard → training/checklists/exams → personal history/competency → notifications/profile"),("Manager","Manager dashboard → employees → training/quality/exam result detail → department reports"),("Admin","Admin dashboard → accounts/reference/settings → training/quality/evaluation configuration → reports"),("Evaluation delegate","Evaluation dashboard → documents/questions/sets/exam management according to granted permissions")],[5,12])
    add_figure(d,"09_navigation.png","Figure 1. Role-based navigation map.","Role-based navigation showing staff, manager, administrator and evaluation-delegate entry points, dashboards and feature routes.")
    d.add_heading("3. Screen Inventory",1)
    for i,r in enumerate(ROUTES,1):
        sid=f"SCR-{i:03d}"; d.add_heading(f"{sid} — {r}",3)
        add_table(d,["Aspect","Design"],[("Access","Protected according to route wrapper and server-side endpoint authorization."),("Components","Page heading/breadcrumb, scoped navigation, data/form area, loading/empty/error feedback."),("Navigation","Entered from role menu, dashboard card or related entity action; unauthorized users are redirected/blocked."),("Data","Calls the corresponding feature API module; list pages use bounded filters/pagination where supported."),("Display conditions","Actions depend on role/permission, entity ownership and lifecycle state."),("Verification",f"Trace to L2 UI/validation and L3/E2E coverage for {sid}.")],[4,13])
    d.add_heading("4. External API Inventory",1); add_table(d,["API-ID","Controller","Base mapping","Methods","Source"],[(f"API-{i:03d}",*row) for i,row in enumerate(CONTROLLERS,1)],[2.5,5,6,2,7])
    d.add_heading("5. Background Job Inventory",1); d.add_heading("5.1 Job Summary Table",2); add_table(d,["JOB-ID","Name","Trigger","Implementation","Outcome"],JOBS,[2,4,4,5,7])
    for job,name,trigger,impl,outcome in JOBS:
        d.add_heading(f"{job} — {name}",2); bullets(d,[f"Trigger: {trigger}.",f"Implementation: {impl}.",f"Processing outcome: {outcome}","Failure behavior: record/log a retryable failure; do not duplicate completed business effects.","Verification: direct method/event execution in L2 and end-state observation in L3 where environment dependencies are available."])
    save(d,out)


def build_testdoc():
    d,out=new_doc("Report 5.0_TestDoc_Template.docx","Report 5.0_TestDoc_CareHub_v0.9.docx","Test Document","Report 5.0")
    d.add_heading("1. Scope of Testing",1); d.add_heading("1.1 In Scope - Functional Requirements",2); p(d,"All 17 FT groups, 34 use cases, 12 business flows, 15 global rules, 59 controller groups, declared routes and 8 background jobs are within the traceability baseline.")
    d.add_heading("1.2 In Scope - Non-Functional Requirements",2); add_table(d,["NFR","Category","Verification"],[(x[0],x[1],x[3]) for x in NFRS],[3,4,10])
    add_figure(d,"10_test_levels.png","Figure 1. Verification levels and current release gates.","Layered verification model showing L1 unit, L2 integration, L3 system and L4 UAT evidence with the remaining release gates.")
    d.add_heading("1.3 Out of Scope",2); bullets(d,["Clinical correctness or medical-device certification.","Hospital production infrastructure not supplied to the team.","Native mobile applications and external SMS/Zalo gateways.","Formal acceptance before Product Owner UAT signature."])
    d.add_heading("1.4 Constraints",2); p(d,"RabbitMQ, PostgreSQL, R2 credentials, external AI keys/models and seeded role accounts are required for complete environment-level testing. Local automated tests may replace selected dependencies with controlled doubles.")
    d.add_heading("1.5 Assumptions",2); p(d,"Test data is synthetic, isolated and resettable. No shared production database is used. Current code behavior overrides stale statements in prior reports.")
    d.add_heading("2. Test Strategy",1); d.add_heading("2.1 Testing Types",2); add_table(d,["Type","Purpose"],[("Functional","Verify UC main/alternative flows and GBRs."),("UI","Verify routes, content states and role-dependent actions."),("Input validation","Verify client-visible and server-side enforcement."),("Background jobs","Verify trigger-independent service behavior, retry and idempotency."),("Performance","Measure normal-load percentiles and scoring/dashboard behavior."),("Security","Verify authentication, authorization, data scope, object access and secret handling.")],[5,12])
    d.add_heading("2.2 Test Levels",2); add_table(d,["Level","Scope","Workbook / tool","Current evidence"],[ ("L1 Unit","Pure/domain/service logic","Report 5.1; JUnit/Vitest","Legacy register: 209 Pass, 8 Fail; current frontend run: 116/116 Pass"),("L2 Integration","Controller/service/repository, UI validation, security, jobs","Report 5.2; Spring Boot Test","Legacy register: 88 Pass, 4 Fail, 1 Blocked"),("L3 System","HTTP/API flows, system security/performance","Report 5.3","Legacy register: 111 Pass, 3 Not Run, 1 Blocked"),("L4 UAT","Business outcome acceptance","Report 5.4","12 scripts prepared; not executed / not signed")],[3,5,6,7])
    d.add_heading("2.3 Supporting Tools",2); p(d,"JUnit 5, Spring Boot Test/MockMvc/real test server, H2/PostgreSQL fixtures, Vitest, Testing Library, Playwright, k6, JaCoCo, Vite build and manual UAT.")
    d.add_heading("3. Test Plan",1); d.add_heading("3.1 Test Environment",2); add_table(d,["Layer","Baseline"],[("Backend","Java 17 target, Spring Boot 4.0.6, test profile"),("Frontend","Node 20+, React 19/Vite 8"),("Data","Isolated H2 or PostgreSQL 17 fixtures"),("Async","RabbitMQ 3 when queue behavior is under test"),("Browser","Current Chrome/Edge; Playwright profiles"),("Performance","k6 against dedicated staging, not developer shared DB")],[5,12])
    d.add_heading("3.2 Test Data",2); p(d,"Minimum data: active/locked users for each role; two departments; activity types and training records in each lifecycle state; evidence objects; a versioned quality form with assignments/submissions; question categories/questions/exam paper/attempts; notification policy/template; AI document job fixtures.")
    d.add_heading("3.3 Test Milestones",2); add_table(d,["Milestone","Exit evidence"],[("M1 L1","Core branch rule target and all Critical cases triaged"),("M2 L2","Controller/security/job integration register executed"),("M3 L3","Critical business flows and NFR spot checks executed"),("M4 UAT","All Critical/High scripts Pass or accepted; Product Owner signs")],[5,12])
    d.add_heading("4. Test Cases",1); d.add_heading("4.1 Test Case Workbooks",2); add_table(d,["Workbook","Purpose","Cases"],[("Report 5.1 Unit Tests L1","Method/domain coverage",217),("Report 5.2 Integration Tests L2","Integration/security/job coverage",93),("Report 5.3 System Tests L3","System/API and performance coverage",115),("Report 5.4 UAT Scripts","Business acceptance",12)],[7,8,2])
    d.add_heading("4.2 Requirements Coverage Matrix",2); add_table(d,["FT-ID","UC","BF","L1","L2","L3","UAT"],[(x[0],x[4],x[5],"As applicable","Required","Required","Mapped by BF") for x in FEATURES],[2,3,3,3,3,3,4])
    d.add_heading("5. Test Reports",1); d.add_heading("5.1 Summary Table",2); add_table(d,["Evidence","Pass","Fail","Skipped/Blocked","Not run","Release interpretation"],[("Current backend suite",794,0,9,0,"BUILD SUCCESS on 11/08/2026; RabbitMQ warnings because the broker was not running"),("Current frontend unit",116,0,0,0,"Observed 11/08/2026"),("Current frontend build",1,0,0,0,"Succeeded with bundle-size warning"),("Legacy L1 register",209,8,0,0,"Historical status; v0.9 execution status starts Not Run"),("Legacy L2 register",88,4,1,0,"Historical status retained"),("Legacy L3 register",111,0,1,3,"Historical status retained"),("UAT",0,0,0,12,"Pending execution/sign-off")],[4,2,2,3,2,7])
    d.add_heading("5.2 Defect Register",2); add_table(d,["ID","Area","Status / required action"],[("DEF-DOC-01","Historical requirement/code divergence","Corrected in v0.9 baseline; verify affected tests."),("DEF-ENV-01","RabbitMQ unavailable during local backend run","Environment dependency; rerun queue tests with broker."),("DEF-PERF-01","SPA bundle >500 kB warning","Open optimization; measure before release."),("DEF-UAT-01","No signed UAT evidence","Release blocker for formal acceptance.")],[3,6,9])
    save(d,out)


def build_deployment():
    d,out=new_doc("Report 6.1_DeploymentGuide_Template.docx","Report 6.1_DeploymentGuide_CareHub_v0.9.docx","Deployment Guide","Report 6.1")
    d.add_heading("Deployment Topology",1)
    add_figure(d,"11_deployment.png","Figure 1. Production-oriented CareHub deployment topology.","CareHub deployment topology connecting user devices, TLS reverse proxy, frontend, backend, core infrastructure and external services.")
    sections=[
        ("1. Prerequisites",[("Software","Linux/Windows server; Java 17+; Node 20+ for build; PostgreSQL 17; RabbitMQ 3; Redis 7; reverse proxy; TLS."),("Hardware","Pilot baseline: 4 CPU, 8 GB RAM, SSD storage sized for DB and object transfer cache."),("Network","Inbound 443; internal backend 8081; DB 5432, AMQP 5672 and Redis 6379 restricted to trusted hosts.")]),
        ("2. Server Preparation",[("OS/runtime","Patch OS, create least-privilege service user, install Java and container/runtime tools."),("Directories","Separate application, configuration, logs, backups and model files; restrict secret file permissions."),("DNS/TLS","Create approved hostname and certificate before exposing the UI.")]),
        ("3. Database Setup",[("Create","Create dedicated carehub database/user; do not use docker-compose default credentials in production."),("Schema","Back up first; validate entity/SQL baseline on staging. Current ddl-auto=update must be replaced by reviewed migration policy for production."),("Verify","Connect with application user and test create/read/update in staging.")]),
        ("4. Application Deployment",[("Backend build","Run mvnw.cmd clean package after tests; deploy the generated JAR with environment configuration."),("Frontend build","Set production VITE_API_BASE_URL, run npm ci and npm run build; publish dist through the reverse proxy."),("Integrity","Record commit hash, artifact hash and deployment date.")]),
        ("5. Reverse Proxy Configuration",[("Routing","Serve SPA static assets; fallback routes to index.html; proxy /api/v1 to backend:8081."),("TLS","Redirect HTTP to HTTPS; set secure headers and upload limits consistent with 20 MB multipart configuration."),("Protection","Never expose storage/model/config/log directories as static paths.")]),
        ("6. Process Manager Service",[("Environment","Load JWT, DB, mail, RabbitMQ, R2, seed and AI variables from a restricted file or secret store."),("Service","Run backend under a dedicated account; restart on failure with bounded retry; capture stdout/stderr."),("Health","Check actuator health and application logs before routing traffic.")]),
        ("7. File Storage Setup",[("R2","Create bucket, least-privilege keys, endpoint and CORS/policy required by server-side access."),("Models","Place E5/VietQuill assets under configured model paths when those providers are enabled."),("Retention","Back up DB metadata and apply object lifecycle only after retention approval.")]),
        ("8. Environment Variables",[("Required","JWT_SECRET, DB_*, ADMIN_*, MAIL_*, RABBITMQ_*, R2_*, GENERATION_API_KEY."),("Optional/tunable","SERVER_PORT, APP_*, GENERATION_*, E5_*, VIETQUILL_*, DOCUMENT_* and validation thresholds."),("Rule","Never use VITE_* for secrets; browser bundles expose them.")]),
        ("9. First-Run Checklist",[("Infrastructure","DB, broker, Redis, SMTP/R2/AI connectivity verified."),("Seed","Run only approved seeds; change the initial admin password immediately."),("Security","TLS, CORS, role guards and evidence download authorization verified."),("Evidence","Record smoke-test results and artifact/commit identifiers.")]),
        ("10. Smoke Test",[("Access","HTTPS redirect; login for ADMIN/MANAGER/USER; protected route rejection."),("Core","Create DRAFT training record; load dashboard; open quality and evaluation lists."),("Async","Publish a test email/job and observe final state; verify retry visibility."),("Files","Upload/download allowed evidence; direct object path remains inaccessible.")]),
        ("11. Backup & Recovery",[("Database","Daily pg_dump plus tested retention and encrypted off-host copy."),("Objects/models","Back up required object metadata/config; model binaries can be restored from approved source."),("Restore","Restore into isolated environment, run integrity queries and smoke tests before failover.")]),
        ("12. Updating the Application",[("Standard","Backup → deploy immutable version → apply reviewed migration → restart → smoke test → monitor."),("Rollback","Stop traffic, restore compatible application/database pair, verify health; never assume schema rollback is automatic."),("Record","Update deployment log with version, owner, result and incident/rollback details.")]),
        ("13. Troubleshooting",[("Startup","Check missing env values, DB reachability, port conflicts and model paths."),("502","Check backend service/port and proxy upstream."),("Login","Check clock, JWT secret consistency, account state and CORS."),("Upload","Check multipart limit, MIME/extension, R2 permission and object cleanup logs."),("Queue","Check RabbitMQ host/credentials/queue declarations and dead-letter/retry state.")]),
    ]
    for title,rows in sections:
        d.add_heading(title,1); add_table(d,["Topic","Procedure / control"],rows,[4.5,13])
    d.add_heading("Appendix A — Ports Summary",1); add_table(d,["Port","Service","Exposure"],[(443,"HTTPS reverse proxy","Public/approved network"),(8081,"Spring Boot","Internal only"),(5432,"PostgreSQL","Internal only"),(5672,"RabbitMQ AMQP","Internal only"),(15672,"RabbitMQ management","Administrator network only"),(6379,"Redis","Internal only")],[3,6,8])
    save(d,out)


def build_manual():
    d,out=new_doc("Report 6.2_UserManual_Template.docx","Report 6.2_UserManual_CareHub_v0.9.docx","User Manual","Report 6.2")
    d.add_heading("About This Manual",1); p(d,"This manual is for Viet Duc Hospital staff, department managers, administrators and delegated evaluation users. Menu visibility depends on assigned roles and permissions.")
    d.add_heading("Things to Know Before You Start",1); bullets(d,["Use a current Chrome or Edge browser and the official HTTPS address.","Do not share passwords, OTPs, access links or downloaded staff evidence.","Save drafts before leaving long forms.","If a page denies access, contact the administrator; hiding a menu is not the security control.","Report errors with time, action and correlation ID—never include passwords or tokens."])
    add_figure(d,"12_user_journey.png","Figure 1. Common user journey across CareHub roles.","Five-step CareHub user journey from sign-in and dashboard through completing work, reviewing outcomes and following up.")
    d.add_heading("Chapter 1 — Getting Started",1)
    for h,text in [("1.1 Supported Browsers","Current Chrome and Edge are the baseline. Mobile-width layouts are available for core staff flows."),("1.2 Accessing CareHub","Open the hospital-provided HTTPS URL. Do not use an IP address or unofficial link in production."),("1.3 Account Provisioning","Accounts are created/imported by an administrator. Self-registration is not available."),("1.4 Signing In","Enter employee code and password. First-time users complete the email/OTP setup flow before normal access."),("1.5 If You Cannot Sign In","Check employee code, password and network. Use Forgot Password when available; contact an administrator for a locked/inactive account."),("1.6 Signing Out","Open the account menu and select Sign out. Close shared-browser windows afterward."),("1.7 Changing Your Password","Open Profile, choose Change Password, enter current/new password and confirm. Never reuse the initial password.")]: d.add_heading(h,2); p(d,text)
    d.add_heading("Chapter 2 — Staff Guide",1)
    add_table(d,["Task","Steps","Outcome"],[("Record training","Training → New → choose activity → enter dates/hours → save draft → attach evidence → submit.","A submitted record appears in personal history and compliance totals when valid."),("Complete checklist","My Checklists → open assignment → answer required items → review → submit.","Submission is scored against the assigned form version."),("Take exam","Exams → select available exam → start → answer → submit before timer expires.","Result and score/10 appear in history when grading completes."),("Review competency","Competency → personal summary/field/technique views.","Employee sees current classification and contributing results."),("Notifications/profile","Open bell/profile menu.","Read notices and maintain permitted personal details.")],[4,8,6])
    d.add_heading("Chapter 3 — Manager Guide",1); add_table(d,["Task","Steps","Control"],[("Monitor team","Manager Dashboard / Employees → filter/select employee.","Only allowed department scope is shown."),("Review training","Training Employees → record/evidence detail → perform permitted action.","State and audit rules are enforced server-side."),("Quality evaluation","Quality Checklists → assignment → evaluate/submit.","Required questions and scoring configuration apply."),("Exam results","Exam Results → employee attempt detail.","Score remains on 0–10 scale."),("Reports","Manager Reports → training/quality/competency dashboard.","Filters remain department scoped.")],[4,8,6])
    d.add_heading("Chapter 4 — Evaluation Delegate Guide",1); add_table(d,["Area","Typical actions"],[("Question documents","Upload document, start generation and monitor job."),("Candidate review","Review source excerpt, duplicate/quality signals; accept or reject."),("Question bank/sets","Create and curate questions/categories/sets within permission."),("Exam management","Configure exam, paper and assignments if EXAM_PUBLISHER is granted."),("Audit/dashboard","Review evaluation activity and outcomes permitted to the account.")],[6,12])
    d.add_heading("Chapter 5 — Administrator Guide",1); add_table(d,["Area","Typical actions","Safety note"],[("Accounts","Create/import, assign role, lock/unlock, reset password.","Verify identity; never communicate credentials in the test record."),("Reference data","Manage departments, positions, education and professional fields.","Changes affect filters and scope across modules."),("System settings","Update global training/compliance configuration.","Record approval and verify dashboard totals after change."),("Quality/evaluation","Manage forms, scoring, targets, question/exam configuration.","Version/publish carefully; recalculation may be asynchronous."),("Notifications","Manage policies/templates and monitor delivery.","Test templates with non-sensitive recipients first."),("Audit/import logs","Review history and rejected rows.","Logs are evidence; do not edit or conceal failures.")],[4,7,7])
    d.add_heading("Frequently Asked Questions",1); add_table(d,["Question","Answer"],[("Why can I not see a menu?","Your account lacks the required role/permission or first-login setup is incomplete."),("Why is my training total unchanged?","Only valid SUBMITTED records in the configured window contribute."),("Can I edit a submitted training record?","No. Only DRAFT is editable; an administrator may return an eligible submitted record to draft."),("Why is an email delayed?","Delivery is queued. Check in-app notification/job state or ask an administrator to inspect RabbitMQ/email logs."),("Can AI questions be used immediately?","No. Generated candidates require authorized human review before entering the bank."),("What should I include in a support request?","Time, page, action, visible message and correlation ID; never send password, OTP or token.")],[7,11])
    save(d,out)


def build_final():
    d,out=new_doc("Report 7_FinalReport_DraftTemplate.docx","Report 7_FinalReport_CareHub_v0.9.docx","Project Handover Report (Draft)","Report 7")
    d.add_heading("1. Executive Summary",1); d.add_heading("1.1 What Was Built",2); p(d,"CareHub is a working web platform that brings staff training, care-quality checklists, competency assessment, notifications and management reporting into one controlled system for Viet Duc Hospital.")
    d.add_heading("1.2 Delivery Status",2); add_table(d,["Area","Status","Evidence / next gate"],[("Functional baseline","Implemented","17 feature groups; current code inventory"),("Frontend unit/build","Pass","116 tests passed; production build succeeded on 11/08/2026"),("Backend automated suite","Pass","794 tests passed, 0 failures/errors, 9 skipped; Maven BUILD SUCCESS on 11/08/2026"),("Messaging environment","Partial","RabbitMQ was not running; the suite completed but queue connectivity produced warnings"),("System/performance","Partial","Legacy registers exist; controlled staging rerun required"),("UAT","Pending","12 business scripts prepared; no signature"),("Production deployment","Pending customer IT","Use Deployment Guide")],[5,4,9])
    add_figure(d,"13_traceability.png","Figure 1. Delivery traceability chain and release evidence gates.","Traceability chain from features through use cases, business flows, design and test levels to UAT and release acceptance.")
    d.add_heading("2. Project Overview",1); d.add_heading("2.1 Background",2); p(d,"The project replaces fragmented spreadsheets, folders and disconnected assessment processes with shared identity, workflow, audit and analytics capabilities.")
    d.add_heading("2.2 Objectives",2); add_table(d,["Objective","Outcome"],[("Trace training compliance","Implemented record/evidence/compliance flows"),("Standardize quality evaluation","Implemented versioned forms, scoring and dashboards"),("Manage competency assessment","Implemented questions, exams, attempts and classification"),("Improve management visibility","Implemented role-scoped dashboards and exports"),("Preserve traceability","Implemented logs/events and rebuilt v0.9 documentation package")],[7,11])
    d.add_heading("2.3 Stakeholders",2); p(d,"Product Owner/supervisor: Nguyễn Văn An. Delivery team: Đỗ Mạnh Tuấn, Dương Ngọc Hiếu, Ngô Trung Kiên, Phí Hoàng Nam and Phạm Hải Nam. Operational users: hospital staff, department managers, administrators and evaluation delegates.")
    d.add_heading("2.4 Project Timeline",2); p(d,"The v0.9 package records the planned 14-week lifecycle and the evidence state observed on 11/08/2026. Historical WBS estimates are retained as planning data, not reported as actual effort.")
    d.add_heading("3. Scope Delivered",1); d.add_heading("3.1 Business Flows",2); add_table(d,["BF-ID","Flow","Status"],[(x[0],x[1],"Implemented; UAT pending") for x in BFS],[2.5,10,5])
    d.add_heading("3.2 Functional Features",2); add_table(d,["FT-ID","Feature","Priority","Status"],[(x[0],x[1],x[2],"Implemented / regression evidence tracked") for x in FEATURES],[2,9,3,5])
    d.add_heading("3.3 Global Business Rules",2); add_table(d,["GBR-ID","Rule","Verification"],[(x[0],x[1],"L1–L3 as applicable") for x in GBRS],[2.5,12,4])
    d.add_heading("3.4 Non-Functional Requirements",2); add_table(d,["NFR","Target","Result"],[(x[0],x[2],x[4]) for x in NFRS],[3,10,6])
    d.add_heading("3.5 Screens / Pages Delivered",2); p(d,f"The router contains {len(ROUTES)} unique declared path entries in this baseline. FDS assigns SCR-001 onward and distinguishes implemented pages from redirect aliases.")
    d.add_heading("3.6 Background Jobs / Scheduled Tasks",2); add_table(d,["JOB-ID","Name","Status"],[(x[0],x[1],"Implemented; environment verification tracked") for x in JOBS],[3,9,6])
    d.add_heading("4. Scope Deferred to Next Version",1); add_table(d,["Item","Reason","Priority"],[("Direct hospital system integration","Requires approved API/data governance","High"),("Native mobile apps","Responsive web is current target","Low"),("SMS/Zalo gateway","External procurement/integration","Medium"),("Production-grade schema migrations","Current ddl-auto baseline must be hardened","High"),("SPA route code-splitting","Build succeeds but bundle warning remains","Medium"),("Formal acceptance","UAT and deployment sign-off pending","Critical gate")],[7,8,3])
    d.add_heading("5. Technical Summary",1); d.add_heading("5.1 Technology Stack",2); p(d,"Java 17 / Spring Boot 4.0.6 backend; React 19 / Vite 8 frontend; PostgreSQL 17; RabbitMQ 3; Redis 7; Cloudflare R2; DeepSeek, E5 and VietQuill integrations.")
    d.add_heading("5.2 Architecture Decisions",2); p(d,"Modular monolith, REST SPA, stateless JWT with role/permission authorization, versioned persistence, event-driven slow work and external adapter boundaries.")
    d.add_heading("5.3 Database Schema",2); p(d,f"The current source contains {len(ENTITIES)} JPA entity classes and 69 repository interfaces observed during the backend test startup. Hibernate update is convenient for development but must be replaced by reviewed versioned migrations for production.")
    d.add_heading("5.4 Security Architecture",2); p(d,"JWT, method authorization, department/user scope, protected object serving, optimistic locking, secret externalization and correlation-aware error handling form the main controls.")
    d.add_heading("5.5 Repository Structure",2); p(d,"carehub-backend contains Java domain packages and infrastructure configuration; carehub-frontend contains app, features and shared JavaScript/CSS; docs contains templates, legacy evidence and this v0.9 CareHub package.")
    d.add_heading("6. Test & Quality Summary",1); d.add_heading("6.1 Test Coverage Summary",2); add_table(d,["Level","Result","Interpretation"],[("Backend automated suite","794 Pass / 0 Fail / 9 Skipped","Current Maven BUILD SUCCESS; broker connectivity warnings"),("Frontend unit","116/116 Pass","Current observed run"),("Frontend build","Pass","Bundle-size warning remains"),("Legacy L1","209 Pass / 8 Fail","Historical workbook status; v0.9 execution status starts Not Run"),("Legacy L2","88 Pass / 4 Fail / 1 Blocked","Historical workbook status"),("Legacy L3","111 Pass / 3 Not Run / 1 Blocked","Historical workbook status"),("UAT","0/12 executed","Formal acceptance pending")],[4,6,8])
    d.add_heading("6.2 Defects Summary",2); p(d,"No claim of zero defects is made. Open gates are tracked in TestDoc: backend/environment rerun, controlled performance evidence, SPA bundle optimization decision and UAT sign-off.")
    d.add_heading("6.4 UAT Scripts",2); p(d,"Twelve UAT scripts map one-to-one to BF-01 through BF-12. Result and signature cells are intentionally blank/Not Executed until a business representative performs them.")
    d.add_heading("7. Delivery Package",1); add_table(d,["Category","Contents"],[("Source","carehub-backend and carehub-frontend at recorded commit"),("SDLC documents","Reports 1, 2.0, 2.1, 3.0, 3.1, 4.0, 4.1, 5.0, 6.1, 6.2 and 7"),("Test artifacts","L1, L2, L3 and UAT workbooks"),("Infrastructure","docker-compose baseline, application/env examples and deployment guide")],[5,13])
    d.add_heading("8. Known Issues & Limitations",1); d.add_heading("8.1 Open Defects Accepted for Release",2); p(d,"None are accepted at this draft stage. Any open Major defect requires explicit Product Owner acceptance before handover.")
    d.add_heading("8.2 Known Limitations",2); bullets(d,["Formal UAT is not executed or signed.","Production-like RabbitMQ/R2/AI and performance evidence is incomplete.","The frontend build reports a large main bundle; code-splitting should be evaluated.","Hibernate ddl-auto update is not the recommended production migration strategy.","Direct hospital-system integration and native mobile clients are outside v0.9 scope."])
    save(d,out)


def main():
    OUT.mkdir(parents=True, exist_ok=True); ARTIFACTS.mkdir(parents=True, exist_ok=True)
    for fn in [build_r1,build_r2,build_prd,build_ucs,build_tds,build_fds,build_testdoc,build_deployment,build_manual,build_final]: fn()
    manifest={"version":VERSION,"date":DATE,"baseline":{"branch":"ManhTuan","commit":"627b9448"},"counts":{"features":len(FEATURES),"use_cases":len(UCS),"business_flows":len(BFS),"routes":len(ROUTES),"controllers":len(CONTROLLERS),"entities":len(ENTITIES),"jobs":len(JOBS)},"documents":[p.name for p in sorted(OUT.glob("*.docx"))]}
    (OUT/"baseline_manifest.json").write_text(json.dumps(manifest,ensure_ascii=False,indent=2),encoding="utf-8")
    canonical={"features":FEATURES,"use_cases":UCS,"business_flows":BFS,"business_rules":GBRS,"nfrs":NFRS,"jobs":JOBS,"routes":ROUTES,"controllers":CONTROLLERS,"entities":ENTITIES}
    (ROOT/"report-work"/"canonical_baseline.json").write_text(json.dumps(canonical,ensure_ascii=False,indent=2),encoding="utf-8")


if __name__ == "__main__": main()
