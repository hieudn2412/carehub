from __future__ import annotations

import json
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from pypdf import PdfReader

ROOT = Path(r"D:\ĐỒ ÁN\carehub")
OUT = ROOT / "docs" / "26FA_v0.9_CareHub"
RENDER = ROOT / "report-work" / "doc-renders"
CANON = json.loads((ROOT / "report-work" / "canonical_baseline.json").read_text(encoding="utf-8"))

DOCS = [
    "Report 1_VS_CareHub_v0.9.docx",
    "Report 2.0_ProjectPlan_CareHub_v0.9.docx",
    "Report 3.0_PRD_CareHub_v0.9.docx",
    "Report 3.1_UCS_CareHub_v0.9.docx",
    "Report 4.0_TDS_CareHub_v0.9.docx",
    "Report 4.1_FDS_CareHub_v0.9.docx",
    "Report 5.0_TestDoc_CareHub_v0.9.docx",
    "Report 6.1_DeploymentGuide_CareHub_v0.9.docx",
    "Report 6.2_UserManual_CareHub_v0.9.docx",
    "Report 7_FinalReport_CareHub_v0.9.docx",
]
BOOKS = [
    "Report 2.1_ProjectTracking_CareHub_v0.9.xlsx",
    "Report 5.1_UnitTests_CareHub_v0.9.xlsx",
    "Report 5.2_IntegrationTests_CareHub_v0.9.xlsx",
    "Report 5.3_SystemTests_CareHub_v0.9.xlsx",
    "Report 5.4_UAT_Scripts_CareHub_v0.9.xlsx",
]


def doc_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        chunks: list[str] = []
        for name in archive.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            root = ET.fromstring(archive.read(name))
            chunks.extend(node.text or "" for node in root.iter() if node.tag.endswith("}t"))
    return " ".join(chunks)


errors: list[str] = []
expected = set(DOCS + BOOKS + ["baseline_manifest.json"])
present = {p.name for p in OUT.iterdir() if p.is_file() and not p.name.endswith(".inspect.ndjson")}
if present != expected:
    errors.append(f"package files mismatch: missing={sorted(expected-present)}, extra={sorted(present-expected)}")

texts: dict[str, str] = {}
page_counts: dict[str, int] = {}
image_counts: dict[str, int] = {}
for name in DOCS:
    path = OUT / name
    try:
        doc = Document(path)
        if not doc.paragraphs and not doc.tables:
            errors.append(f"empty document: {name}")
        text = doc_text(path)
        texts[name] = text
        image_counts[name] = len(doc.inline_shapes)
        if not doc.inline_shapes:
            errors.append(f"no diagram in {name}")
        for shape in doc.inline_shapes:
            if not shape._inline.docPr.get("descr"):
                errors.append(f"missing image alt text in {name}")
        for marker in ["{{", "[Project Name]", "TalentHub", "Order Fulfillment", "Lorem ipsum"]:
            if marker.casefold() in text.casefold():
                errors.append(f"template remnant {marker!r}: {name}")
        pdf = RENDER / f"{path.stem}.pdf"
        page_counts[name] = len(PdfReader(pdf).pages)
    except Exception as exc:
        errors.append(f"cannot parse {name}: {exc}")

checks = {
    "Report 3.0_PRD_CareHub_v0.9.docx": [
        *[x[0] for x in CANON["features"]],
        *[x[0] for x in CANON["use_cases"]],
        *[x[0] for x in CANON["nfrs"]],
    ],
    "Report 3.1_UCS_CareHub_v0.9.docx": [
        *[x[0] for x in CANON["use_cases"]],
        *[x[0] for x in CANON["business_flows"]],
    ],
    "Report 4.0_TDS_CareHub_v0.9.docx": [
        *[x[0] for x in CANON["nfrs"] if x[1] in {"Performance", "Reliability"}],
    ],
    "Report 4.1_FDS_CareHub_v0.9.docx": [
        *[x[0] for x in CANON["jobs"]],
    ],
    "Report 5.0_TestDoc_CareHub_v0.9.docx": [x[0] for x in CANON["nfrs"]],
}
for name, ids in checks.items():
    missing = [item for item in ids if item not in texts.get(name, "")]
    if missing:
        errors.append(f"missing traceability ids in {name}: {missing}")

print(json.dumps({
    "documents": len(DOCS),
    "workbooks": len(BOOKS),
    "pages": page_counts,
    "total_pages": sum(page_counts.values()),
    "images": image_counts,
    "total_images": sum(image_counts.values()),
    "errors": errors,
}, ensure_ascii=False, indent=2))
raise SystemExit(1 if errors else 0)
